"""
Local Document Service

Handles document upload, storage, and text extraction without external dependencies.
Encrypts document content at rest for data protection.
"""

import logging
import os
import uuid
from pathlib import Path

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import Document
from app.services.encryption import get_encryption_service

logger = logging.getLogger(__name__)

# Document storage directory
UPLOAD_DIR = Path(settings.UPLOAD_DIR if hasattr(settings, 'UPLOAD_DIR') else "./uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


async def extract_text(file_content: bytes, mime_type: str, filename: str) -> str:
    """
    Extract text content from various file types.

    Supports: PDF, DOCX, TXT, MD, HTML
    """
    text = ""

    try:
        if mime_type == "text/plain" or filename.endswith(('.txt', '.md')):
            # Plain text / Markdown
            text = file_content.decode('utf-8', errors='ignore')

        elif mime_type == "application/pdf" or filename.endswith('.pdf'):
            # PDF extraction
            try:
                from io import BytesIO

                import PyPDF2

                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except ImportError:
                logger.warning("PyPDF2 not installed, PDF text extraction unavailable")
                text = "[PDF content - install PyPDF2 for text extraction]"
            except Exception as e:
                logger.error(f"PDF extraction error: {e}")
                text = f"[PDF extraction failed: {e!s}]"

        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or filename.endswith('.docx'):
            # DOCX extraction
            try:
                from io import BytesIO

                from docx import Document as DocxDocument

                doc = DocxDocument(BytesIO(file_content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except ImportError:
                logger.warning("python-docx not installed, DOCX text extraction unavailable")
                text = "[DOCX content - install python-docx for text extraction]"
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                text = f"[DOCX extraction failed: {e!s}]"

        elif mime_type == "text/html" or filename.endswith('.html'):
            # HTML - strip tags for basic text
            try:
                from html.parser import HTMLParser

                class TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text_parts = []

                    def handle_data(self, data):
                        self.text_parts.append(data)

                parser = TextExtractor()
                parser.feed(file_content.decode('utf-8', errors='ignore'))
                text = ' '.join(parser.text_parts)
            except Exception as e:
                logger.error(f"HTML extraction error: {e}")
                text = file_content.decode('utf-8', errors='ignore')
        else:
            # Try to decode as text
            try:
                text = file_content.decode('utf-8', errors='ignore')
            except Exception:
                text = "[Binary content - text extraction not supported]"

    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        text = f"[Text extraction failed: {e!s}]"

    return text.strip()


async def save_file(
    file_content: bytes,
    filename: str,
    organization_id: str
) -> str:
    """
    Save file to local filesystem.

    Returns the file path.
    """
    # Create organization directory
    org_dir = UPLOAD_DIR / organization_id
    org_dir.mkdir(parents=True, exist_ok=True)

    # Generate unique filename
    file_ext = Path(filename).suffix
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = org_dir / unique_filename

    # Write file
    with open(file_path, 'wb') as f:
        f.write(file_content)

    return str(file_path)


async def upload_document(
    db: AsyncSession,
    file_content: bytes,
    filename: str,
    mime_type: str,
    organization_id: str,
    user_id: str,
    title: str | None = None,
    author: str | None = None,
    description: str | None = None
) -> Document:
    """
    Upload and process a document.

    1. Save file to filesystem
    2. Extract text content
    3. Encrypt content for storage
    4. Store metadata in database
    """
    # Save file
    file_path = await save_file(file_content, filename, organization_id)

    # Extract text
    text_content = await extract_text(file_content, mime_type, filename)

    # Encrypt content for storage
    encryption = get_encryption_service()
    encrypted_content = encryption.encrypt_document_content(text_content)

    # Create document record
    document = Document(
        organization_id=organization_id,
        title=title or filename,
        author=author,
        description=description,
        filename=filename,
        file_path=file_path,
        file_size=len(file_content),
        mime_type=mime_type,
        content=encrypted_content,
        source_system="upload",
        status="uploaded"
    )

    db.add(document)
    await db.commit()
    await db.refresh(document)

    logger.info(f"Document uploaded: {document.document_id} - {filename} (encrypted: {encryption.is_enabled})")

    return document


async def get_documents(
    db: AsyncSession,
    organization_id: str,
    page: int = 1,
    page_size: int = 20,
    search: str | None = None
) -> tuple[list[Document], int]:
    """
    Get paginated list of documents for an organization.
    """
    # Base query
    query = select(Document).where(Document.organization_id == organization_id)

    # Search filter
    if search:
        search_filter = f"%{search}%"
        query = query.where(
            (Document.title.ilike(search_filter)) |
            (Document.filename.ilike(search_filter)) |
            (Document.author.ilike(search_filter))
        )

    # Order by created_at desc
    query = query.order_by(Document.created_at.desc())

    # Get total count
    count_query = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar() or 0

    # Apply pagination
    offset = (page - 1) * page_size
    query = query.offset(offset).limit(page_size)

    # Execute
    result = await db.execute(query)
    documents = result.scalars().all()

    return list(documents), total


async def get_document(
    db: AsyncSession,
    document_id: str,
    organization_id: str
) -> Document | None:
    """
    Get a single document by ID.
    """
    result = await db.execute(
        select(Document).where(
            Document.document_id == document_id,
            Document.organization_id == organization_id
        )
    )
    return result.scalar_one_or_none()


async def delete_document(
    db: AsyncSession,
    document_id: str,
    organization_id: str
) -> bool:
    """
    Delete a document and its file.
    """
    document = await get_document(db, document_id, organization_id)

    if not document:
        return False

    # Delete file if exists
    if document.file_path and os.path.exists(document.file_path):
        try:
            os.remove(document.file_path)
        except Exception as e:
            logger.error(f"Error deleting file {document.file_path}: {e}")

    # Delete database record
    await db.delete(document)
    await db.commit()

    logger.info(f"Document deleted: {document_id}")

    return True


async def get_document_stats(
    db: AsyncSession,
    organization_id: str
) -> dict:
    """
    Get document statistics for an organization.
    """
    # Total documents
    total_query = select(func.count()).select_from(Document).where(
        Document.organization_id == organization_id
    )
    total_result = await db.execute(total_query)
    total_documents = total_result.scalar() or 0

    # Total size
    size_query = select(func.sum(Document.file_size)).where(
        Document.organization_id == organization_id
    )
    size_result = await db.execute(size_query)
    total_size = size_result.scalar() or 0

    # Documents by status
    status_query = select(
        Document.status,
        func.count(Document.document_id)
    ).where(
        Document.organization_id == organization_id
    ).group_by(Document.status)
    status_result = await db.execute(status_query)
    status_counts = {row[0]: row[1] for row in status_result.fetchall()}

    return {
        "total_documents": total_documents,
        "total_size_bytes": total_size,
        "status_counts": status_counts
    }


def get_decrypted_content(document: Document) -> str:
    """
    Get decrypted content from a document.

    Args:
        document: Document model instance

    Returns:
        Decrypted content string
    """
    if not document.content:
        return ""

    encryption = get_encryption_service()
    return encryption.decrypt_document_content(document.content)
