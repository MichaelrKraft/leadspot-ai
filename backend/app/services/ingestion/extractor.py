"""
Content Extraction Service

Extracts text content from various document formats:
- PDF (PyMuPDF)
- Microsoft Office (docx, xlsx, pptx)
- Google Docs/Sheets
- HTML/Markdown
- Plain text
"""

import io
import logging
from pathlib import Path
from typing import Any

import aiofiles
import fitz  # PyMuPDF
import markdown
import openpyxl
from bs4 import BeautifulSoup
from docx import Document

logger = logging.getLogger(__name__)


class ContentExtractor:
    """Service for extracting content from various document formats."""

    # Supported MIME types
    SUPPORTED_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # docx
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # xlsx
        'application/msword',  # doc (legacy)
        'text/html',
        'text/markdown',
        'text/plain',
        'application/vnd.google-apps.document',
        'application/vnd.google-apps.spreadsheet',
    }

    async def extract(
        self,
        file_path: str | None = None,
        file_content: bytes | None = None,
        mime_type: str | None = None,
        source_url: str | None = None
    ) -> dict[str, Any]:
        """
        Extract content from a document.

        Args:
            file_path: Path to local file
            file_content: Raw file bytes
            mime_type: MIME type of the file
            source_url: URL of the source (for Google Docs, web pages)

        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata (title, author, pages, etc.)
                - format: Source format
                - success: Boolean indicating success
                - error: Error message if failed

        Raises:
            ValueError: If neither file_path nor file_content provided
        """
        if not file_path and not file_content:
            raise ValueError("Either file_path or file_content must be provided")

        try:
            # Read file content if path provided
            if file_path and not file_content:
                async with aiofiles.open(file_path, 'rb') as f:
                    file_content = await f.read()

            # Detect MIME type if not provided
            if not mime_type:
                mime_type = self._detect_mime_type(file_path, file_content)

            logger.info(f"Extracting content from {mime_type}")

            # Route to appropriate extractor
            if mime_type == 'application/pdf':
                result = await self._extract_pdf(file_content)
            elif 'wordprocessingml' in mime_type or mime_type == 'application/msword':
                result = await self._extract_docx(file_content)
            elif 'spreadsheetml' in mime_type:
                result = await self._extract_xlsx(file_content)
            elif mime_type == 'text/html':
                result = await self._extract_html(file_content)
            elif mime_type == 'text/markdown':
                result = await self._extract_markdown(file_content)
            elif mime_type == 'text/plain':
                result = await self._extract_text(file_content)
            elif 'google-apps.document' in mime_type:
                result = await self._extract_google_doc(source_url)
            elif 'google-apps.spreadsheet' in mime_type:
                result = await self._extract_google_sheet(source_url)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {mime_type}',
                    'text': '',
                    'metadata': {},
                    'format': mime_type
                }

            result['success'] = True
            result['format'] = mime_type
            return result

        except Exception as e:
            logger.error(f"Error extracting content: {e!s}", exc_info=True)
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {},
                'format': mime_type or 'unknown'
            }

    def _detect_mime_type(
        self,
        file_path: str | None,
        file_content: bytes
    ) -> str:
        """Detect MIME type from file extension or content."""
        if file_path:
            extension = Path(file_path).suffix.lower()
            type_map = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.html': 'text/html',
                '.md': 'text/markdown',
                '.txt': 'text/plain',
            }
            if extension in type_map:
                return type_map[extension]

        # Default to text/plain
        return 'text/plain'

    async def _extract_pdf(self, content: bytes) -> dict[str, Any]:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(stream=content, filetype="pdf")

        text_parts = []
        metadata = {
            'pages': len(doc),
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
            'keywords': doc.metadata.get('keywords', ''),
        }

        for page_num, page in enumerate(doc, 1):
            # Extract text with layout preservation
            text = page.get_text("text")
            if text.strip():
                # Add page marker
                text_parts.append(f"\n--- Page {page_num} ---\n{text}")

        doc.close()

        return {
            'text': '\n'.join(text_parts),
            'metadata': metadata
        }

    async def _extract_docx(self, content: bytes) -> dict[str, Any]:
        """Extract text from Word document."""
        doc = Document(io.BytesIO(content))

        text_parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
        }

        # Try to get core properties
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
            })
        except Exception:
            pass

        return {
            'text': '\n\n'.join(text_parts),
            'metadata': metadata
        }

    async def _extract_xlsx(self, content: bytes) -> dict[str, Any]:
        """Extract text from Excel spreadsheet."""
        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)

        text_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

            for row in sheet.iter_rows(values_only=True):
                # Convert row to strings and filter empty cells
                row_text = ' | '.join(
                    str(cell) for cell in row
                    if cell is not None and str(cell).strip()
                )
                if row_text.strip():
                    text_parts.append(row_text)

        metadata = {
            'sheets': len(workbook.sheetnames),
            'sheet_names': workbook.sheetnames,
        }

        workbook.close()

        return {
            'text': '\n'.join(text_parts),
            'metadata': metadata
        }

    async def _extract_html(self, content: bytes) -> dict[str, Any]:
        """Extract text from HTML."""
        soup = BeautifulSoup(content, 'html.parser')

        # Remove script and style elements
        for script in soup(['script', 'style', 'noscript']):
            script.decompose()

        # Get text
        text = soup.get_text(separator='\n', strip=True)

        # Extract metadata
        metadata = {
            'title': soup.title.string if soup.title else '',
        }

        # Meta tags
        meta_description = soup.find('meta', attrs={'name': 'description'})
        if meta_description:
            metadata['description'] = meta_description.get('content', '')

        return {
            'text': text,
            'metadata': metadata
        }

    async def _extract_markdown(self, content: bytes) -> dict[str, Any]:
        """Extract text from Markdown."""
        text = content.decode('utf-8', errors='ignore')

        # Convert to HTML then extract text for cleaner output
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        clean_text = soup.get_text(separator='\n', strip=True)

        # Extract title (first H1)
        title = ''
        h1 = soup.find('h1')
        if h1:
            title = h1.get_text(strip=True)

        return {
            'text': clean_text,
            'metadata': {'title': title}
        }

    async def _extract_text(self, content: bytes) -> dict[str, Any]:
        """Extract plain text."""
        text = content.decode('utf-8', errors='ignore')

        return {
            'text': text,
            'metadata': {}
        }

    async def _extract_google_doc(self, url: str) -> dict[str, Any]:
        """
        Extract text from Google Doc.

        Note: This requires Google Drive API setup.
        For MVP, this is a placeholder.
        """
        # TODO: Implement Google Drive API integration
        logger.warning("Google Docs extraction not yet implemented")
        return {
            'text': '',
            'metadata': {},
            'error': 'Google Docs extraction requires Drive API setup'
        }

    async def _extract_google_sheet(self, url: str) -> dict[str, Any]:
        """
        Extract text from Google Sheet.

        Note: This requires Google Drive API setup.
        For MVP, this is a placeholder.
        """
        # TODO: Implement Google Drive API integration
        logger.warning("Google Sheets extraction not yet implemented")
        return {
            'text': '',
            'metadata': {},
            'error': 'Google Sheets extraction requires Drive API setup'
        }
